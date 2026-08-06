#!/usr/bin/env python3

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    # Since the template doesn't have standard heading styles, we'll create our own
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add spacing after heading
    paragraph.paragraph_format.space_after = Pt(12)
    return paragraph

def add_subheading(doc, text):
    # For subheadings, we'll use a slightly smaller font size
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add spacing after subheading
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph

def add_paragraph(doc, text, bold=False, font_size=10.5):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    # Add spacing after paragraph
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph

def main():
    if len(sys.argv) != 3:
        print("Usage: {} <template_docx> <output_docx>".format(sys.argv[0]))
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        # Load the template document
        doc = Document(template_path)
        
        # Clear existing content (optional, depending on your needs)
        # For now, we'll just append to the template
        
        # Basic Information Section
        add_heading(doc, "基本信息")
        add_paragraph(doc, "姓名：张艳江")
        add_paragraph(doc, "籍贯：河北-保定")
        add_paragraph(doc, "工作经验：10年")
        add_paragraph(doc, "邮箱：814643107@qq.com")
        add_paragraph(doc, "政治面貌：中共党员")
        add_paragraph(doc, "电话：13521622479")
        
        # Job Objective Section
        add_heading(doc, "求职意向")
        add_paragraph(doc, "职位：嵌入式Linux开发工程师")
        add_paragraph(doc, "期望薪资：面议")
        add_paragraph(doc, "到岗时间：随时到岗")
        
        # Skills Section
        add_heading(doc, "技能特长")
        skills = [
            "熟悉 C/C++ 开发编程，了解QT图形页面化编程，LVGL（GUI）开发，良好的编码风格；",
            "熟练使用 linux 操作系统，熟练使用gcc、g++、gdb等GUN工具集；",
            "熟悉常用的数据结构及算法；",
            "ARM(cortex-A) 开发经验；",
            "熟悉 ARM常用接口技术(UART、IIC、SPI)；",
            "熟悉ARM常用外设开发(GPIO、PWM)；",
            "熟悉 linux 移植相关工作(u-boot、kernel、root)；",
            "熟练掌握linux 字符设备驱动开发；",
            "熟悉linux 各类(GPIO、IIC、中断输入)子系统驱动开发；",
            "熟练掌握 platform平台总线驱动开发；",
            "熟悉linux 内核竞态解决方法及中断处理机制；",
            "熟悉V4L2摄像头驱动框架开发；",
            "熟悉socket网络编程；",
            "熟练掌握LinuxIO、多进程/线程开发，熟悉进程间通讯机制及线程同步互斥机制；",
            "了解块设备驱动开发；"
        ]
        for skill in skills:
            add_paragraph(doc, "• " + skill)
        
        # Work Experience Section
        add_heading(doc, "工作/项目经验")
        
        # Experience 1
        add_subheading(doc, "2024-10 ~ 至今  北京G7物联")
        add_paragraph(doc, "职位：嵌入式Linux开发工程师")
        add_paragraph(doc, "项目：商用车(货车)行车Ai智驾安全设备（紫宝盒--AiBox)")
        add_paragraph(doc, "开发语言：C，C++")
        add_paragraph(doc, "开发平台：Ubuntu，Windows")
        add_paragraph(doc, "开发工具：VScode")
        add_paragraph(doc, "项目介绍：")
        add_paragraph(doc, "• 智能硬件瑞芯微rv1126 32位系统，后续rv1126b 64位系统。")
        add_paragraph(doc, "• 远程司机打卡，小7云对讲，云平台数据管理。")
        add_paragraph(doc, "• 8路摄像头实时视频远程平台查看。")
        add_paragraph(doc, "个人职责：")
        add_paragraph(doc, "• 主要实现了8路USB摄像头对接Ai算法模型，实时监控行车安全报警，QT实时输出报警。")
        add_paragraph(doc, "• 主要负责Ai算法的对接业务模块，8路摄像头的算法对接配置，以及给算法喂帧图片，作为报警业务生产者；")
        add_paragraph(doc, "• 算法有：ADAS，DMS，BSD，目标检测，环境分类检测，安全帽检测；")
        add_paragraph(doc, "• 8路摄像头的视频编码H264；")
        add_paragraph(doc, "• 8路定时拍照，RGA的硬件转换nv12转jpg，以及优化编码的不同转换；")
        add_paragraph(doc, "• 图片上的OSD的水印添加，LVGL的编写；")
        add_paragraph(doc, "• 报警事件的图片的保存，水印添加；")
        add_paragraph(doc, "• 进程间通信zmq，数据结构jsoncpp，图像opencv，等等第三方库的集成；")
        add_paragraph(doc, "• 日志系统glog，物联通信MQTT等等业务；")
        
        # Experience 2
        add_subheading(doc, "2023-06 ~ 2024-09  北京至简科技有限公司")
        add_paragraph(doc, "项目：车机手机互联盒子（KALOS2、AuTo Pro X)")
        add_paragraph(doc, "开发语言：C")
        add_paragraph(doc, "开发平台：Ubuntu")
        add_paragraph(doc, "开发工具：VScode")
        add_paragraph(doc, "项目介绍：")
        add_paragraph(doc, "• 智能硬件瑞芯微rv1109，和全志v851S3，处理芯片。")
        add_paragraph(doc, "• 主要实现了AndroidAuto协议，通过手机和盒子蓝牙互联然后热点连接，")
        add_paragraph(doc, "• 盒子与车机USB连接，从而由有线转无线连接车机。")
        add_paragraph(doc, "• 可以手机车机自动连接，车技实时显示手机的功能APP，可以导航播放音乐等。")
        add_paragraph(doc, "• 可以手机dex连接，实现dex手机投屏到车机。")
        add_paragraph(doc, "个人职责：")
        add_paragraph(doc, "• Auto协议的更改修复bug，协议版本的不同添加；")
        add_paragraph(doc, "• LVGL新功能的的更改、开发；")
        add_paragraph(doc, "• 自有产品KALOS2的视频编码h264的优化，以及其他bug的修复优化；")
        add_paragraph(doc, "• 韩国项目AutoProX的对接维护开发，出差韩国对接不同车型遇到的问题，查找修复；")
        
        # Experience 3
        add_subheading(doc, "2020-04 ~ 2023-05  北京美沃斯科技文化有限公司")
        add_paragraph(doc, "职位：嵌入式工程师")
        add_paragraph(doc, "项目：智能医美室云系统(cortex-A7)")
        add_paragraph(doc, "开发语言：C、html")
        add_paragraph(doc, "开发平台：Ubuntu、windows")
        add_paragraph(doc, "开发工具：VScode")
        add_paragraph(doc, "项目介绍：")
        add_paragraph(doc, "• 项目基于什么C/S模型+B/S模型实现，分为三部分，智能硬件(MP157A)、云系统、用户终端。")
        add_paragraph(doc, "• 主要实现了用户通过链接云系统实现对智能硬件所在场景的环境实时监控、环境数据采集、阈值设置、设备控制等功能。")
        add_paragraph(doc, "• 云系统对用户信息采用sqlite3数据库管理。系统支持并发处理。用户终端为手机或PC，通过浏览器链接服务器，实现登录进入系统。")
        add_paragraph(doc, "• 服务器采用 BOA web服务器实现用户网页交互功能，云系统采用 cookie 技术实现用户ID的识别功能。")
        add_paragraph(doc, "应用技术：")
        add_paragraph(doc, "• boa 服务器移植、mjpg-streamer视频流服务器移植、多进程编程、多线程编程、进程间通讯(msgqueueu)、线程同步互斥(无名信号量)、cgi开发、sqlite3数据库开、　　TCP并发服务器开发、智能硬件客户端app开发。")
        add_paragraph(doc, "• u-boot移植、linux内核移植、rootfs文件系统制作。")
        add_paragraph(doc, "• Linux字符设备驱动开发(GPIO\\IIC\\PWM\\SPI\\按键中断)。")
        add_paragraph(doc, "• ov5640摄像头驱动配置。")
        add_paragraph(doc, "个人职责：")
        add_paragraph(doc, "• 项目框架设计；")
        add_paragraph(doc, "• mjpg-streamer视频流服务器移植；")
        add_paragraph(doc, "• TCP并发服务器开发及各功能线程开发；")
        add_paragraph(doc, "• MP157A TCP客户端开发及各功能线程开发；")
        add_paragraph(doc, "• u-boot移植、linux内核移植、rootfs文件系统制作；")
        add_paragraph(doc, "• Linux字符设备驱动开发(GPIO\\IIC\\PWM\\SPI\\按键中断)；")
        add_paragraph(doc, "• ov5640摄像头驱动配置；")
        
        # Additional Projects
        add_subheading(doc, "美沃心电检测仪器")
        add_paragraph(doc, "开发语言：C，C++")
        add_paragraph(doc, "开发工具：Ubuntu、windows")
        add_paragraph(doc, "开发工具：VScode、qt_creator")
        add_paragraph(doc, "项目介绍：")
        add_paragraph(doc, "• 智能STMMP157a的A核硬件SOC，ad8232心率传感器获取心率数据。")
        add_paragraph(doc, "• 主要实现了通过心率传感器经过ADC转换，由串口输出显示心率波形。")
        add_paragraph(doc, "• 可以配置采样率，采样时长，由定时器中断控制。")
        add_paragraph(doc, "个人职责：")
        add_paragraph(doc, "• 项目框架设计；")
        add_paragraph(doc, "• u-boot移植、linux内核移植、rootfs文件系统制作；")
        add_paragraph(doc, "• Linux字符设备驱动开发(UART\\ADC\\中断处理)；")
        add_paragraph(doc, "• qt通讯协议；")
        
        # More Experiences
        add_subheading(doc, "2016-04 ~ 2020-03  时机科技北京有限公司")
        add_paragraph(doc, "项目：时机智能语音助手")
        add_paragraph(doc, "开发语言：C")
        add_paragraph(doc, "开发平台：Ubuntu")
        add_paragraph(doc, "开发工具：VScode")
        add_paragraph(doc, "职位：软件开发工程师")
        add_paragraph(doc, "项目介绍：")
        add_paragraph(doc, "• 项目利用百度平台api接口的语音识别（rest-api-asr）、语音合成（reset-api-tts），结合硬件进行语音控制系统；")
        add_paragraph(doc, "• 项目主要用STM32MP157A arm的板子进行开发；")
        add_paragraph(doc, "• 当用户说出关键词：时机时机。进行语音交互，进行控制硬件操作。")
        add_paragraph(doc, "• 设置关键词：打开灯光，进行控制硬件操作；打开电视，控制电视开关电源，打开窗帘，控制窗帘的关闭。")
        add_paragraph(doc, "• uboot移植，tf-a移植，sqlite3移植，Linux内核移植，进程间通信，脚本文件的编写；")
        add_paragraph(doc, "• 简易设置闹钟功能：定时器闹钟。")
        add_paragraph(doc, "应用技术：")
        add_paragraph(doc, "• rootfs根文件系统；")
        add_paragraph(doc, "• 录制的pcm文件转wav文件；")
        add_paragraph(doc, "• Linux字符设备驱动开发(GPIO\\IIC\\中断)。")
        add_paragraph(doc, "个人职责：")
        add_paragraph(doc, "• 利用busybox制作rootfs根文件系统；")
        add_paragraph(doc, "• Linux字符设备驱动开发(GPIO\\IIC\\PWM\\SPI\\中断)。")
        add_paragraph(doc, "• uboot移植，tf-a移植，sqlite3移植，Linux内核移植，QT移植；")
        
        # Education Background
        add_heading(doc, "教育背景")
        add_paragraph(doc, "学校：河北工业职业技术学院")
        
        # Self Evaluation
        add_heading(doc, "自我评价")
        add_paragraph(doc, "• 工作积极认真，细心负责，熟练运用办公自动化软件，善于在工作中提出问题、发现问题、解决问题，有较强的分析能力;")
        add_paragraph(doc, "• 勤奋好学，踏实肯干，动手能力强，认真负责，有很强的社会责任感;")
        add_paragraph(doc, "• 坚毅不拔，吃苦耐劳，喜欢迎接新挑战，为人和善，善于沟通。")
        
        # Save the document
        doc.save(output_path)
        print(f"Resume generated successfully to {output_path}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()